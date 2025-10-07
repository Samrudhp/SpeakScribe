"""
Export service for generating PDF and DOCX reports
"""
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import logging
from typing import Dict
from pathlib import Path

logger = logging.getLogger(__name__)


class ExportService:
    def __init__(self):
        """Initialize export service"""
        self.export_dir = Path("exports")
        self.export_dir.mkdir(exist_ok=True)
    
    def generate_pdf(self, data: Dict) -> str:
        """
        Generate PDF report
        
        Args:
            data: Meeting data dictionary
            
        Returns:
            Path to generated PDF
        """
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Title
            pdf.set_font("Arial", "B", 20)
            pdf.cell(0, 10, "Meeting Summary Report", ln=True, align="C")
            pdf.ln(5)
            
            # Date
            pdf.set_font("Arial", "I", 10)
            pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
            pdf.ln(5)
            
            # Summary Section
            if data.get('summary'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Executive Summary", ln=True)
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 6, data['summary'])
                pdf.ln(5)
            
            # Sentiment Section
            if data.get('sentiment'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Overall Sentiment", ln=True)
                pdf.set_font("Arial", "", 11)
                sentiment = data['sentiment']
                pdf.cell(0, 6, f"Tone: {sentiment.get('overall_sentiment', 'N/A')} {sentiment.get('overall_emoji', '')}", ln=True)
                pdf.cell(0, 6, f"Emotional Tone: {sentiment.get('emotional_tone', 'N/A')}", ln=True)
                pdf.ln(5)
            
            # Action Items
            if data.get('actions') and data['actions'].get('action_items'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Action Items", ln=True)
                pdf.set_font("Arial", "", 11)
                for item in data['actions']['action_items']:
                    pdf.multi_cell(0, 6, f"• {item}")
                pdf.ln(5)
            
            # Decisions
            if data.get('actions') and data['actions'].get('decisions'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Decisions Made", ln=True)
                pdf.set_font("Arial", "", 11)
                for item in data['actions']['decisions']:
                    pdf.multi_cell(0, 6, f"• {item}")
                pdf.ln(5)
            
            # Topics
            if data.get('topics'):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Key Topics", ln=True)
                pdf.set_font("Arial", "", 11)
                for topic in data['topics']:
                    pdf.cell(0, 6, f"• {topic}", ln=True)
                pdf.ln(5)
            
            # Transcript
            if data.get('transcript'):
                pdf.add_page()
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Full Transcript", ln=True)
                pdf.set_font("Arial", "", 10)
                pdf.multi_cell(0, 5, data['transcript'])
            
            # Save PDF
            filename = f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filepath = self.export_dir / filename
            pdf.output(str(filepath))
            
            logger.info(f"PDF generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"PDF generation error: {str(e)}")
            raise
    
    def generate_docx(self, data: Dict) -> str:
        """
        Generate DOCX report
        
        Args:
            data: Meeting data dictionary
            
        Returns:
            Path to generated DOCX
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading("Meeting Summary Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].italic = True
            
            doc.add_paragraph()
            
            # Summary Section
            if data.get('summary'):
                doc.add_heading("Executive Summary", 1)
                doc.add_paragraph(data['summary'])
            
            # Sentiment Section
            if data.get('sentiment'):
                doc.add_heading("Overall Sentiment", 1)
                sentiment = data['sentiment']
                doc.add_paragraph(
                    f"Tone: {sentiment.get('overall_sentiment', 'N/A')} "
                    f"{sentiment.get('overall_emoji', '')}"
                )
                doc.add_paragraph(
                    f"Emotional Tone: {sentiment.get('emotional_tone', 'N/A')}"
                )
            
            # Action Items
            if data.get('actions') and data['actions'].get('action_items'):
                doc.add_heading("Action Items", 1)
                for item in data['actions']['action_items']:
                    doc.add_paragraph(item, style='List Bullet')
            
            # Decisions
            if data.get('actions') and data['actions'].get('decisions'):
                doc.add_heading("Decisions Made", 1)
                for item in data['actions']['decisions']:
                    doc.add_paragraph(item, style='List Bullet')
            
            # Deadlines
            if data.get('actions') and data['actions'].get('deadlines'):
                doc.add_heading("Deadlines", 1)
                for item in data['actions']['deadlines']:
                    doc.add_paragraph(item, style='List Bullet')
            
            # Topics
            if data.get('topics'):
                doc.add_heading("Key Topics", 1)
                for topic in data['topics']:
                    doc.add_paragraph(topic, style='List Bullet')
            
            # Transcript
            if data.get('transcript'):
                doc.add_page_break()
                doc.add_heading("Full Transcript", 1)
                doc.add_paragraph(data['transcript'])
            
            # Save DOCX
            filename = f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = self.export_dir / filename
            doc.save(str(filepath))
            
            logger.info(f"DOCX generated: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"DOCX generation error: {str(e)}")
            raise


# Singleton instance
_export_service = None

def get_export_service() -> ExportService:
    """Get or create ExportService instance"""
    global _export_service
    if _export_service is None:
        _export_service = ExportService()
    return _export_service
